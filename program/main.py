import os
import sys
import shutil
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import win32com.client as win32

# ===== 공통 설정 =====
APP_TITLE = "재범방지교육통합센터 문서 발급기"
ACCENT_COLOR = "#2C73D2"  # 파란 포인트(B)
BG_COLOR = "#F5F7FB"
TEXT_COLOR = "#222222"

# 템플릿 경로 처리 (templates 폴더 사용)
def base_path(filename: str) -> str:
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, "templates", filename)
    return os.path.abspath(os.path.join("templates", filename))

# 기본 문서 6개 (항상 PDF 생성)
BASE_TEMPLATES = {
    "인지행동개선훈련 증명서.pdf": "template_인지행동개선훈련.docx",
    "준법의식교육 증명서.pdf": "template_준법의식.docx",
    "준법생활 계획서": "template_준법생활 계획서.docx",
    "재범 위험 종합 관리 평가 증명서.pdf": "template_재범 위험 종합 관리 평가 증명서.docx",
    "재범방지교육통합센터 탄원서.pdf": "template_재범방지교육통합센터_탄원서.docx",
    "재범방지교육통합센터 소견서.pdf": "template_재범방지교육통합센터_소견서.docx"
}

# 선택형 템플릿 (큰 문서, 보조 문서, 출력명 앞부분)
SELECT_TEMPLATES = {
    "1": ("template_재범방지교육_재산범죄.docx", "template_재산범죄.docx", "재산범죄"),
    "2": ("template_재범방지교육_도박중독.docx", "template_도박중독.docx", "도박중독"),
    "3": ("template_재범방지교육_디지털범죄.docx", "template_디지털범죄.docx", "디지털범죄"),
    "4": ("template_재범방지교육_마약범죄.docx", "template_마약범죄.docx", "마약범죄"),
    "5": ("template_재범방지교육_성범죄.docx", "template_성범죄.docx", "성범죄"),
    "6": ("template_재범방지교육_음주운전.docx", "template_음주운전.docx", "음주운전"),
    "7": ("template_재범방지교육_폭력범죄.docx", "template_폭력범죄.docx", "폭력범죄")
}

# GUI에 보여줄 선택 옵션
OPTION_LABELS = [
    "선택 안함",
    "재산범죄",
    "도박중독",
    "디지털범죄",
    "마약범죄",
    "성범죄",
    "음주운전",
    "폭력범죄",
]
OPTION_TO_KEY = {
    "선택 안함": "0",
    "재산범죄": "1",
    "도박중독": "2",
    "디지털범죄": "3",
    "마약범죄": "4",
    "성범죄": "5",
    "음주운전": "6",
    "폭력범죄": "7",
}

# DOCX 그대로 복사만 하는 문서 3개
DOCX_COPY_TEMPLATES = {
    "심리상담사 소견서.docx": "template_심리상담사 소견서.docx",
    "변호사 상담증명서.docx": "template_변호사 상담증명서.docx",
    "반성문.docx": "template_반성문.docx"
}

# PDF 변환만 필요한 문서 3개
PDF_ONLY_TEMPLATES = {
    "각 이수 소감문 가이드.pdf": "template_각 이수 소감문 가이드라인, 소감문 예시.docx",
    "심리상담 소감문 가이드.pdf": "template_심리상담 소감문 작성 가이드라인, 소감문 예시.docx",
    "반성문 탄원서 작성 가이드.pdf": "template_반성문, 탄원서 작성 가이드 양식.docx"
}

# ===== DOCX 유틸 =====
def bold_replace(paragraphs, label, new_line):
    for p in paragraphs:
        if p.text.strip().startswith(label):
            p.text = ""
            run = p.add_run(new_line)
            run.bold = True
            return

def replace_text(paragraphs, old, new):
    for p in paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)

def remove_extra(paragraphs, label):
    found = []
    for p in paragraphs:
        if p.text.strip().startswith(label):
            found.append(p)
    for p in found[1:]:
        p._element.getparent().remove(p._element)

def generate_pdf(template_docx, pdf_path, name, date, word, is_plan=False):
    doc = Document(template_docx)

    if is_plan:
        replace_text(doc.paragraphs, "계획서 작성일자:", f"계획서 작성일자: {date}")
        replace_text(doc.paragraphs, "작성자:", f"작성자: {name}")
        replace_text(doc.paragraphs, "작성자 :", f"작성자 : {name}")
    else:
        remove_extra(doc.paragraphs, "발급일자")
        remove_extra(doc.paragraphs, "발급기관")
        bold_replace(doc.paragraphs, "발급일자", f"발급일자: {date}")
        bold_replace(doc.paragraphs, "발급기관", "발급기관: 재범방지교육통합센터")
        bold_replace(doc.paragraphs, "수강생 성명", f"수강생 성명: 재범방지교육통합센터 수강생 {name}")

    temp_docx = pdf_path.replace(".pdf", "_temp.docx")
    doc.save(temp_docx)

    wdoc = word.Documents.Open(temp_docx)
    wdoc.SaveAs(pdf_path, FileFormat=17)
    wdoc.Close()
    os.remove(temp_docx)

def generate_pdf_no_replace(template_docx, pdf_path, word):
    temp_docx = pdf_path.replace(".pdf", "_temp.docx")
    shutil.copy(template_docx, temp_docx)

    wdoc = word.Documents.Open(temp_docx)
    wdoc.SaveAs(pdf_path, FileFormat=17)
    wdoc.Close()
    os.remove(temp_docx)

# ===== 문서 생성 처리 =====
def run_generation(student_name: str, date_str: str, option_label: str):
    option_key = OPTION_TO_KEY.get(option_label, "0")

    # 출력 폴더
    folder = f"재범방지교육통합센터 수강생 {student_name}"
    folder_path = os.path.join(os.path.abspath("."), folder)
    os.makedirs(folder_path, exist_ok=True)

    # Word 실행
    word = win32.Dispatch("Word.Application")
    word.Visible = False

    try:
        # 1) 기본 PDF 생성
        for output_name, tmpl in BASE_TEMPLATES.items():
            if output_name == "준법생활 계획서":
                pdf_name = f"{student_name} 준법생활 계획서.pdf"
                generate_pdf(
                    base_path(tmpl),
                    os.path.join(folder_path, pdf_name),
                    student_name,
                    date_str,
                    word,
                    is_plan=True,
                )
            else:
                generate_pdf(
                    base_path(tmpl),
                    os.path.join(folder_path, output_name),
                    student_name,
                    date_str,
                    word,
                )

        # 2) DOCX 복사 생성 3개
        for out_name, tmpl in DOCX_COPY_TEMPLATES.items():
            new_name = f"{student_name} {out_name}"
            shutil.copy(
                base_path(tmpl),
                os.path.join(folder_path, new_name)
            )

        # 3) PDF only 문서 생성
        for out_name, tmpl in PDF_ONLY_TEMPLATES.items():
            pdf_name = f"{student_name} {out_name}"
            generate_pdf_no_replace(
                base_path(tmpl),
                os.path.join(folder_path, pdf_name),
                word,
            )

        # 4) 선택형 문서 생성
        if option_key != "0":
            big_tmpl, sub_tmpl, prefix = SELECT_TEMPLATES[option_key]

            big_output = os.path.join(
                folder_path,
                "재범방지교육통합센터 교육내용 증명서.pdf",
            )
            generate_pdf(
                base_path(big_tmpl),
                big_output,
                student_name,
                date_str,
                word,
            )

            sub_output = os.path.join(
                folder_path,
                f"{prefix} 재범방지교육 교육내용 증명서.pdf"
            )
            generate_pdf(
                base_path(sub_tmpl),
                sub_output,
                student_name,
                date_str,
                word,
            )

    finally:
        word.Quit()

    # 폴더 자동 열기
    try:
        os.startfile(folder_path)
    except:
        pass

    return folder_path

# ===== GUI =====
def create_gui():
    root = tk.Tk()
    root.title(APP_TITLE)
    root.configure(bg=BG_COLOR)

    root.geometry("520x330")
    root.resizable(False, False)

    main_frame = tk.Frame(root, bg=BG_COLOR)
    main_frame.pack(fill="both", expand=True, padx=24, pady=20)

    title_label = tk.Label(
        main_frame,
        text=APP_TITLE,
        bg=BG_COLOR,
        fg=TEXT_COLOR,
        font=("Malgun Gothic", 16, "bold"),
    )
    title_label.grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))

    subtitle_label = tk.Label(
        main_frame,
        text="수강생 정보를 입력하고 발급할 재범방지교육 유형을 선택하세요.",
        bg=BG_COLOR,
        fg="#555555",
        font=("Malgun Gothic", 10),
    )
    subtitle_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 16))

    # 이름 입력
    name_label = tk.Label(
        main_frame,
        text="수강생 성명",
        bg=BG_COLOR,
        fg=TEXT_COLOR,
        font=("Malgun Gothic", 11),
    )
    name_label.grid(row=2, column=0, sticky="e", pady=5, padx=(0, 10))

    name_var = tk.StringVar()
    name_entry = tk.Entry(
        main_frame,
        textvariable=name_var,
        font=("Malgun Gothic", 11),
        width=25,
        relief="solid",
        bd=1,
    )
    name_entry.grid(row=2, column=1, sticky="w", pady=5)

    # 발급일자 입력
    date_label = tk.Label(
        main_frame,
        text="발급일자",
        bg=BG_COLOR,
        fg=TEXT_COLOR,
        font=("Malgun Gothic", 11),
    )
    date_label.grid(row=3, column=0, sticky="e", pady=5, padx=(0, 10))

    date_var = tk.StringVar()
    date_entry = tk.Entry(
        main_frame,
        textvariable=date_var,
        font=("Malgun Gothic", 11),
        width=25,
        relief="solid",
        bd=1,
    )
    date_entry.grid(row=3, column=1, sticky="w", pady=5)

    # 선택형 콤보박스
    option_label = tk.Label(
        main_frame,
        text="선택형 재범방지교육",
        bg=BG_COLOR,
        fg=TEXT_COLOR,
        font=("Malgun Gothic", 11),
    )
    option_label.grid(row=4, column=0, sticky="e", pady=5, padx=(0, 10))

    option_var = tk.StringVar(value=OPTION_LABELS[0])
    option_combo = ttk.Combobox(
        main_frame,
        textvariable=option_var,
        values=OPTION_LABELS,
        state="readonly",
        font=("Malgun Gothic", 10),
        width=23,
    )
    option_combo.grid(row=4, column=1, sticky="w", pady=5)

    # 상태 메세지
    status_var = tk.StringVar(value="")
    status_label = tk.Label(
        main_frame,
        textvariable=status_var,
        bg=BG_COLOR,
        fg="#777777",
        font=("Malgun Gothic", 9),
    )
    status_label.grid(row=5, column=0, columnspan=2, sticky="w", pady=(10, 0))

    # 버튼 실행 함수
    def on_generate():
        name = name_var.get().strip()
        date_str = date_var.get().strip()
        opt_label = option_var.get().strip()

        if not name:
            messagebox.showerror("입력 오류", "수강생 성명을 입력해주세요.")
            return
        if not date_str:
            messagebox.showerror("입력 오류", "발급일자를 입력해주세요.")
            return

        status_var.set("문서 생성 중입니다. 잠시만 기다려주세요...")
        root.update_idletasks()

        try:
            folder_path = run_generation(name, date_str, opt_label)
        except Exception as e:
            status_var.set("")
            messagebox.showerror("오류", f"문서 생성 중 오류가 발생했습니다.\n\n{e}")
            return

        status_var.set("모든 문서가 생성되었습니다.")
        messagebox.showinfo(
            "완료",
            f"문서 생성이 완료되었습니다.\n\n폴더가 자동으로 열리지 않았다면 아래 경로를 확인하세요.\n\n{folder_path}",
        )

    # 버튼
    generate_btn = tk.Button(
        main_frame,
        text="문서 생성하기",
        command=on_generate,
        bg=ACCENT_COLOR,
        fg="white",
        activebackground="#2457A5",
        activeforeground="white",
        font=("Malgun Gothic", 11, "bold"),
        relief="flat",
        padx=20,
        pady=6,
        cursor="hand2",
    )
    generate_btn.grid(row=6, column=0, columnspan=2, pady=(20, 0))

    root.mainloop()

if __name__ == "__main__":
    create_gui()
